
# Describe:- 
# This code is reading .zip file from the current directory.
# After reading the files it rad each the  files 
# if its pdf, .txt, .docx it will read the documents 
# from those documents it will extract the question's asnwers from the files.

import os
import docx
#import mysql.connector

import zipfile
folder=	zipfile.ZipFile("C:/Users/Prashant/Desktop/checknow.zip", 'r')
file_list_name=folder.namelist()
print(file_list_name)
# for name in file_list_name:
# 	print(name)
num_question=10
num_question_list=[]
text_page_List=[]
dicto=dict()
the_text=[]
for i in range(num_question):
	num_question_list.append("Q"+str(i+1)+")")

import pdflib
# directory=os.listdir("/home/sudhir/Documents/New folder")
# For traverse  in the Directory.
# files_in=folder.open()
for files in file_list_name:
	# print(files)
	text_page_List=[]

	dicto={}
	# print(files)
	index=files.find('.')
	file_name=files[index+1:]
	# print(	directory)
	#If file is text file
	if file_name=='txt':
		# print("TXT")
		file=open(files,'r')
		text_page_List=file.read().splitlines()   # spliting into lines
		# print("The text")
		# print(file_list)


	#If file is pdf file
	if file_name=='pdf':


	# For reading the PDFs.	
		docs=pdflib.Document(files)
		strans=""


		k=""

		# indicate=False  
		page=''
		text_page=''
		found=False
		
		for lines1 in docs:


			

			#Page making  joining all the lines.
			page='\n'.join(lines1.lines).strip()





			#Here appending all the lines into a list
			for pager in page.splitlines():
				text_page_List.append(pager)
				text_page=text_page+" "+pager+'\n'

		# print("The PDF")
		# print(text_page_List) 


	  # The extracted PDF.


	#If file is Docs file
	if file_name=='docx':

		# print("docx")

		readdoc=docx.Document(files)

		for paragraph in readdoc.paragraphs:

			text_page_List.append(paragraph.text)

		# print("The docs")
		# print(text_page_List)	
#for question numbers in the form Q1),

	# dicto=dict()
	
	
	next1=''
	# print(i)
	indicate=False
	# print(text_page_List)
	if (len(text_page_List)>1):
		ans=""
		previous=""
		for line in  text_page_List:

			a1=line.strip()
			if a1 in num_question_list:

					previous=next1
					next1=a1
					# print(store_index)
					found=True
					if  indicate ==True:
						# ans="\n".s)
						# dicto is a dictonary to store all answers.
						dicto[previous]=ans
						# # print(ans)
						# answer_lists[previous]=ans
						previous=next1
						next1=a1
						ans=""
						continue
					elif indicate==False	:
						indicate=True


						ans=""
						continue		
				
				# continue
			if indicate==True:
				# if found:
				# 	found=False
				# 	continue
				# # ans="\n".join(ans)
				ans=ans+" "+line
					# print(asn)
			# breakif
		# For the last question.
			dicto[previous]=ans
		# print(dicto)	
		# print(dicto['Q3)'])

		if ('Q5)' in dicto.keys()):
			print("Q5)")
			print(dicto['Q5)'])
