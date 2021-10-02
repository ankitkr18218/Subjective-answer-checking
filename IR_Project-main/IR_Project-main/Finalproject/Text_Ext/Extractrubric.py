# Describe:- 
# This code is reading .zip file from the current directory.
# After reading the files it rad each the  files 
# if its pdf, .txt, .docx it will read the documents 
# from those documents it will extract the question's asnwers from the files.

class Question:
	question_type=''
	marks_list=[]
	partial=False
	answer=''
	answer_list=	[]

object1=Question()
object1.answer="oewoireiiewre"

# print(object1.answer)
import os
import docx


# import zipfile
# folder=	zipfile.ZipFile("/home/sudhir/Documents/New folder/checknow.zip",'r')
# file_list_name=folder.namelist()
# print(file_list_name)
# # for name in file_list_name:
# 	print(name)		
def runrubric(num_question):
	# num_question=num_question
	num_question_list=[]
	text_page_List=[]
	dicto=dict()
	the_text=[]
	for i in range(num_question):
		num_question_list.append("Q"+str(i+1)+")")


	text_page_List=[]

	files='Rubric_format.docx'          #filepath or filename if in a same directory
	dicto={}
	# print(files)
	index=files.find('.')
	file_name=files[index+1:]
	#If file is text file
	if file_name=='txt':
		# print("TXT")
		file=open(files,'r')
		text_page_List=file.read().splitlines()   # spliting into lines
		

	#If file is pdf file
	if file_name=='pdf':


	# For reading the PDFs.	
		docs=pdflib.Document(files)
		strans=""

		k=""

		page=''
		text_page=''
		found=False
		
		for lines1 in docs:
			
			#Page making  joining all the lines.
			page='\n'.join(lines1.lines)

			#Here appending all the lines into a list
			for pager in page.splitlines():
				text_page_List.append(pager)
				text_page=text_page+" "+pager+'\n'

	  # The extracted PDF.

	#If file is Docs file
	if file_name=='docx':

		readdoc=docx.Document(files)

		for paragraph in readdoc.paragraphs:

			text_page_List.append(paragraph.text)

	next12=''
		# print(i)
	indicate1=False	
	# print(text_page_List)
	if (len(text_page_List)>1):
		ans=""
		previous=""
		question_type=''
		partial=False
		marks_list=''
		marks=[]	
		answer_list=[]

		for line in  text_page_List:
			# if(line=='\n'):
			# 		print('oqeiqip')
			a1=line
			a2=line[0:3]
			if a2 	in num_question_list:

					previous=next12
					next12=a2
					# print(store_index)

					found=True
					if  indicate1 ==True:
						# ans="\n".s)
						# dicto is a dictonary to store all answers.
						# aqq1=ans.strip().lower().replace('.','').replace(',','')
						# aqq1=ans.strip().lower().replace('.'',',' ')
						if previous!='':
							Q22	=Question()
							# print("testtttt",marks_list)
							Q22.question_type=question_type
							Q22.partial=partial
							Q22.marks_list=list(marks)
							Q22.answer_list=answer_list
							Q22.answer=ans
							dicto[previous]=Q22	
							answer_list=[]
							index1=a1.find('[')
							index2=a1.find(']')
							# print(index2)
							question_type =a1[index1+1:index2].lower()
							# print(question_type)
							a3=a1[index2+1:]
							partial=False
							if (question_type=='subjective'):
								index3=a3.find('[')
								index4=a3.find(']')
								# print(index3)	

								marks_list =a3[index3+1:index4].lower()
								# print(marks_list)
								if marks_list.lower()=='partial':
									partial=True
								else:
									partial=False
								a3=a3[index4+1:
]
							index3=a3.find('[')
							index4=a3.find('mark')
							# a3=a3[index3+1: index4	]

							# print(a3	)
							marks =list(map(int,a3[index3+1:index4	].split('+')))
							
						previous=next12
						next12=a2
						ans=""
						continue
					elif indicate1==False	:
						index1=a1.find('[')
						index2=a1.find(']')
						# print(index2)
						question_type =a1[index1+1:index2].lower()
						# print(question_type)
						a3=a1[index2+1:]
						partial=False

						if (question_type=='subjective'):
								index3=a3.find('[')
								index4=a3.find(']')
								# print(index3)	

								marks_list =a3[index3+1:index4].lower()
								# print(marks_list)
								if marks_list.lower()=='partial':
									partial=True
								else:
									partial=False
								a3=a3[index4+1:]
						index3=a3	.find('[')
						index4=a3	.find('mark')

						marks =list(map(int,a3	[index3+1:index4	].split('+')))
						# print(marks)
						indicate1=True


						ans=""
						continue		
				
				# continue
			if indicate1==True:
	
				line.strip()
				if line.strip()!='' :
					# print("line---------",line)
					answer_list.append(line.strip())
				ans=(ans+"\n "+line).strip()
				
			if previous!='':

				Q22	=Question()
				# print(marks_list)
				Q22.question_type=question_type
				Q22.partial=partial
				Q22.marks_list=list(marks)
				Q22.answer=ans
				Q22.answer_list=answer_list

				dicto[previous]=Q22

		#if ('Q5)' in dicto.keys()):
			#print("Q5)")
			#print(len(dicto['Q5)'].answer) )
		return dicto

'''dictoo=runrubric(10)
for i in dictoo.keys():
			print(i)
			print(dictoo[i].answer)
			print(dictoo[i].answer_list)
			print(dictoo[i].marks_list)
			print(dictoo[i].question_type)
			print(dictoo[i].partial)'''
#print(runrubric(10).keys())
